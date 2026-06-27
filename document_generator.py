import os
from datetime import datetime
from typing import Optional

import markdown
from xhtml2pdf import pisa

try:
    from docx import Document
except ImportError:
    Document = None


def ensure_export_folder(path: str = "exportaciones") -> str:
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)
    return path


def generate_word(content: str, filename: str = "documento_mdr.docx", template_path: Optional[str] = None) -> str:
    if Document is None:
        raise ImportError("python-docx no está instalado. Instala python-docx para generar archivos .docx.")

    output_dir = ensure_export_folder()
    output_path = os.path.join(output_dir, filename)

    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    doc.add_paragraph(content)
    doc.save(output_path)
    return output_path


def export_html_to_pdf(html_content: str, output_path: str) -> None:
    with open(output_path, "wb") as pdf_file:
        pisa_status = pisa.CreatePDF(html_content, dest=pdf_file)
    if pisa_status.err:
        raise RuntimeError("Error al generar PDF con xhtml2pdf.")


def generate_pdf_from_markdown(markdown_text: str, filename: str = "documento_mdr.pdf") -> str:
    output_dir = ensure_export_folder()
    output_path = os.path.join(output_dir, filename)

    html_content = markdown.markdown(markdown_text, extensions=["tables", "fenced_code"])
    styled_html = f"""
    <html>
      <head>
        <meta charset='utf-8' />
        <style>
          body {{ font-family: Arial, sans-serif; font-size: 11pt; margin: 2cm; }}
          h1, h2, h3 {{ color: #0f2e5a; }}
          table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
          th, td {{ border: 1px solid #bbb; padding: 6px; }}
          th {{ background-color: #e8f0ff; }}
        </style>
      </head>
      <body>
        {html_content}
      </body>
    </html>
    """

    export_html_to_pdf(styled_html, output_path)
    return output_path


def create_export_filename(prefix: str = "mdr_document") -> str:
    return f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
